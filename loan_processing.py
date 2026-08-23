"""
loan_processing.py — document-set processing + eligibility assessment per loan type.

Backs the Landing Page's "Loan Processing" boxes (Home / Vehicle / Mortgage /
Personal). One box submits a job:

    {loan_type, input_path, output_path, prompt}

and the job runs on its own thread with EXACTLY TWO AGENTS:

    Planning Agent   (CLAUDE_LOAN_PLANNER.md)   one call — reads the document
                     inventory (names only) and the operator prompt, and writes
                     the processing plan the second agent follows.
    Processing Agent (CLAUDE_LOAN_PROCESSOR.md) N+1 calls — one per document
                     (drives the progress bar), then one final call that decides
                     eligibility across the whole file.

Everything is streamed to the UI as SSE events (`doc_started`, `doc_completed`,
`job_completed`, …) carrying running token counts, and written to output_path:

    <output_path>/processing_plan.json     the Planning Agent's plan
    <output_path>/documents/<doc>.json     per-document findings
    <output_path>/eligibility_report.json  structured decision
    <output_path>/eligibility_report.md    readable decision
    <output_path>/eligibility_report.html  printable report (batch-report styling)
    <output_path>/summary.json             job manifest (paths, prompt, tokens)

The LLM provider follows the same config as the orchestrator pipeline
(core.config: DEFAULT_PROVIDER / WORKER_MODEL), so the whole app stays on one
model configuration.

NOTE ON PATHS: input_path/output_path are server-side filesystem paths supplied
by the operator, exactly like the pipeline's `codebase` argument. This is a
single-operator local tool; there is no path sandboxing.
"""
from __future__ import annotations

import hashlib
import json
import queue
import re
import threading
import time
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from core.config import (log, AGENTS_DIR, LLM_PROVIDER, OLLAMA_BASE_URL,
                         WORKER_MODEL)

# ─────────────────────────────────────────────────────────────────────────────
# Loan types + default (operator-editable) prompts
# ─────────────────────────────────────────────────────────────────────────────

_COMMON_TAIL = """
For every document in the set:
  - identify the document type and who it belongs to,
  - pull out the figures and dates that matter for this decision,
  - flag anything missing, expired, inconsistent, or altered.

Then assess eligibility against the criteria above and return one of
ELIGIBLE / NOT_ELIGIBLE / NEEDS_REVIEW, with the specific evidence behind it.
Never invent a value that is not in the documents — say what is missing instead.
"""

DEFAULT_PROMPTS: dict[str, str] = {
    "home": (
        "You are assessing a HOME LOAN application from the supplied document set.\n\n"
        "Eligibility criteria:\n"
        "  - Applicant age 21-65 at maturity; stable income for the last 6 months.\n"
        "  - FOIR / EMI-to-income ratio at or below 50%.\n"
        "  - Loan-to-value at or below 80% of the property's assessed value.\n"
        "  - Credit score 700+, no write-offs or settlements in the last 24 months.\n"
        "  - Clear property title, approved building plan, and valid encumbrance "
        "certificate.\n"
        "  - KYC complete: identity, address, income proof, and bank statements.\n"
        + _COMMON_TAIL
    ),
    "vehicle": (
        "You are assessing a VEHICLE LOAN application from the supplied document set.\n\n"
        "Eligibility criteria:\n"
        "  - Applicant age 21-60; valid driving licence.\n"
        "  - Net monthly income at or above the lender's floor; EMI-to-income at or "
        "below 45%.\n"
        "  - Loan-to-value at or below 85% of ex-showroom price (new) or valuation "
        "(used).\n"
        "  - Credit score 675+, no vehicle-loan default history.\n"
        "  - Proforma invoice / quotation, insurance, and registration details "
        "consistent with the applicant.\n"
        "  - KYC complete: identity, address, income proof, and bank statements.\n"
        + _COMMON_TAIL
    ),
    "mortgage": (
        "You are assessing a MORTGAGE / LOAN-AGAINST-PROPERTY application from the "
        "supplied document set.\n\n"
        "Eligibility criteria:\n"
        "  - Applicant age 25-65 at maturity; verifiable income or business vintage "
        "of 3+ years.\n"
        "  - Loan-to-value at or below 65% of the assessed market value.\n"
        "  - FOIR at or below 55%; existing obligations fully disclosed.\n"
        "  - Credit score 700+, no ongoing litigation on the property.\n"
        "  - Title deed chain, latest tax receipts, valuation report, and "
        "encumbrance certificate present and mutually consistent.\n"
        "  - KYC complete: identity, address, income proof, and bank statements.\n"
        + _COMMON_TAIL
    ),
    "personal": (
        "You are assessing a PERSONAL LOAN application from the supplied document "
        "set.\n\n"
        "Eligibility criteria:\n"
        "  - Applicant age 23-58; salaried with 1+ year total work experience or "
        "self-employed with 2+ years vintage.\n"
        "  - Net monthly income at or above the lender's floor; EMI-to-income at or "
        "below 40%.\n"
        "  - Credit score 720+, no delinquency over 30 days in the last 12 months.\n"
        "  - Salary credits in the bank statement matching the declared income.\n"
        "  - No more than 3 active unsecured loans.\n"
        "  - KYC complete: identity, address, income proof, and bank statements.\n"
        + _COMMON_TAIL
    ),
}

LOAN_TYPES: list[dict[str, str]] = [
    {"id": "home",     "label": "Home Loan",     "icon": "H", "domain": "loan"},
    {"id": "vehicle",  "label": "Vehicle Loan",  "icon": "V", "domain": "loan"},
    {"id": "mortgage", "label": "Mortgage Loan", "icon": "M", "domain": "loan"},
    {"id": "personal", "label": "Personal Loan", "icon": "P", "domain": "loan"},
]

# ── Account processing ───────────────────────────────────────────────────────
# Same engine, different document sets. CSV statement exports are parsed and
# totalled in code (loan_extractors.parse_csv_statement), so the statement box
# starts on parse-first; ID and address documents have no parser, so KYC and
# the general box start on ai_first. Either way an unmatched layout is sent to
# the model in full.
ACCOUNT_TYPES: list[dict[str, str]] = [
    {"id": "account_statement", "label": "Account Processing", "icon": "A",
     "domain": "account", "default_mode": "deterministic"},
    {"id": "kyc",               "label": "KYC Processing",     "icon": "K",
     "domain": "account", "default_mode": "ai_first"},
    {"id": "general",           "label": "General",            "icon": "G",
     "domain": "account", "default_mode": "ai_first"},
]

PROCESSING_TYPES: list[dict[str, str]] = LOAN_TYPES + ACCOUNT_TYPES


def type_spec(type_id: str) -> dict | None:
    return next((t for t in PROCESSING_TYPES if t["id"] == type_id), None)


DEFAULT_PROMPTS.update({
    "account_statement": (
        "You are reviewing a CURRENT/SAVINGS ACCOUNT for the statement period "
        "covered by the supplied documents.\n\n"
        "Report, for the period:\n"
        "  - Total debits: the sum of every debit, and the count of debit "
        "transactions.\n"
        "  - Total credits: the sum of every credit, and the count of credit "
        "transactions.\n"
        "  - Average balance: the mean of the daily closing balances across the "
        "period. State whether it is a daily average or a month-end average, and "
        "say which days were carried forward when no transaction occurred.\n"
        "  - Cheques returned: the number of cheques returned or dishonoured, "
        "with the date, amount, and stated reason for each.\n\n"
        "Also flag anything that would concern a reviewer: a balance chain that "
        "does not reconcile against the stated totals, an opening or closing "
        "balance that disagrees with the transactions, round-tripping, or "
        "returns clustered near period end.\n\n"
        "Every figure must come from the documents. Where the statement does not "
        "cover the whole period, say so rather than extrapolating; where a figure "
        "cannot be computed from what is present, mark it unverified and state "
        "what is missing."
    ),
    "kyc": (
        "You are performing a KYC completeness check on the supplied document "
        "set.\n\n"
        "Establish for the customer:\n"
        "  - Identity proof: which document, its number (masked to the last four "
        "characters), issuing authority, and validity dates.\n"
        "  - Address proof: which document, the address as printed, and whether "
        "it is current.\n"
        "  - Photograph and signature: present or absent.\n"
        "  - PAN / tax identifier, and whether it is legible and consistent.\n"
        "  - Date of birth, and consistency of name and DOB across every "
        "document.\n\n"
        "Then report:\n"
        "  - Documents present, documents missing, and documents expired or "
        "expiring within 90 days.\n"
        "  - Any mismatch in name spelling, address, DOB, or identifier between "
        "documents.\n"
        "  - Signs of tampering: altered figures, inconsistent fonts, or a "
        "scanned page that does not match the rest of the set.\n\n"
        "Conclude with KYC COMPLETE, KYC INCOMPLETE, or NEEDS REVIEW, and list "
        "exactly what must be collected to close any gap. Never treat an absent "
        "document as satisfied."
    ),
    "general": (
        "Review the supplied documents and report what they contain.\n\n"
        "Replace this prompt with your own instructions — this box exists for "
        "ad-hoc work that does not fit the standard checks. State what you want "
        "extracted, the criteria to apply, and the shape of the answer you want "
        "back.\n\n"
        "Until then: identify each document, pull out the figures, dates and "
        "parties that matter, note anything missing, inconsistent or expired, and "
        "say plainly what a reviewer should look at first."
    ),
})

# ─────────────────────────────────────────────────────────────────────────────
# Document discovery + text extraction
# ─────────────────────────────────────────────────────────────────────────────

TEXT_SUFFIXES  = {".txt", ".md", ".json", ".csv", ".log", ".eml", ".htm", ".html"}
PDF_SUFFIXES   = {".pdf"}
DOCX_SUFFIXES  = {".docx"}
EXCEL_SUFFIXES = {".xlsx", ".xls"}
SUPPORTED = TEXT_SUFFIXES | PDF_SUFFIXES | DOCX_SUFFIXES | EXCEL_SUFFIXES

SKIP_DIRS = {".git", ".venv", "venv", "node_modules", "__pycache__", ".idea", ".vscode"}

MAX_DOCS      = 200      # per job
MAX_DOC_CHARS = 60_000   # per document sent to the model
MAX_PASS2_CHARS = 120_000


def scan_documents(root: Path) -> tuple[list[Path], list[Path]]:
    """Return (supported, skipped) files under root, sorted, recursive."""
    supported: list[Path] = []
    skipped: list[Path] = []
    if root.is_file():
        (supported if root.suffix.lower() in SUPPORTED else skipped).append(root)
        return supported, skipped
    for p in sorted(root.rglob("*")):
        if not p.is_file() or p.name.startswith("."):
            continue
        if any(part in SKIP_DIRS for part in p.parts):
            continue
        (supported if p.suffix.lower() in SUPPORTED else skipped).append(p)
    return supported[:MAX_DOCS], skipped


def extract_text(path: Path) -> tuple[str, str]:
    """Return (text, error). Exactly one of the two is non-empty."""
    suffix = path.suffix.lower()
    try:
        if suffix in TEXT_SUFFIXES:
            return path.read_text(encoding="utf-8", errors="replace"), ""

        if suffix in PDF_SUFFIXES:
            try:
                import pdfplumber
            except ImportError:
                try:
                    from pypdf import PdfReader          # optional fallback
                except ImportError:
                    return "", ("no PDF reader installed — "
                                "pip install pdfplumber (or pypdf)")
                reader = PdfReader(str(path))
                return "\n\n".join((pg.extract_text() or "") for pg in reader.pages), ""
            with pdfplumber.open(str(path)) as pdf:
                pages = [(pg.extract_text() or "") for pg in pdf.pages]
            text = "\n\n".join(pages).strip()
            if not text:
                return "", "PDF has no extractable text layer (scanned image?)"
            return text, ""

        if suffix in DOCX_SUFFIXES:
            try:
                import docx                              # python-docx
            except ImportError:
                return "", "python-docx is not installed — pip install python-docx"
            doc = docx.Document(str(path))
            body = "\n".join(p.text for p in doc.paragraphs)
            tables = [
                "\t".join(c.text for c in row.cells)
                for t in doc.tables for row in t.rows
            ]
            return "\n".join([body, *tables]).strip(), ""

        if suffix in EXCEL_SUFFIXES:
            try:
                import pandas as pd
            except ImportError:
                return "", "pandas is not installed — pip install pandas openpyxl"
            sheets = pd.read_excel(str(path), sheet_name=None)
            return "\n\n".join(
                f"--- sheet: {name} ---\n{df.to_csv(index=False)}"
                for name, df in sheets.items()
            ), ""

        return "", f"unsupported file type: {suffix or '(none)'}"
    except Exception as exc:                                          # noqa: BLE001
        return "", f"{type(exc).__name__}: {exc}"


def _truncate(text: str, limit: int) -> tuple[str, bool]:
    if len(text) <= limit:
        return text, False
    return text[:limit], True


# ─────────────────────────────────────────────────────────────────────────────
# LLM plumbing (same provider config as the orchestrator pipeline)
# ─────────────────────────────────────────────────────────────────────────────

class LoanConfigError(RuntimeError):
    """Raised for operator-fixable setup problems (missing key, bad path)."""


def _make_llm(max_tokens: int):
    import os
    if LLM_PROVIDER == "ollama":
        from langchain_ollama import ChatOllama
        return ChatOllama(model=WORKER_MODEL, base_url=OLLAMA_BASE_URL)
    key = os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        raise LoanConfigError(
            "ANTHROPIC_API_KEY is not set — add it to .env (or set "
            "DEFAULT_PROVIDER=ollama to run against a local model)."
        )
    from langchain_anthropic import ChatAnthropic
    return ChatAnthropic(model=WORKER_MODEL, anthropic_api_key=key,
                         max_tokens=max_tokens)


class TruncatedReplyError(RuntimeError):
    """The model hit its output cap mid-answer, so the reply is unusable."""


def _invoke(llm, system: str, user: str) -> tuple[str, int, int]:
    """Return (text, input_tokens, output_tokens).

    A reply cut off at the output cap is raised rather than returned: it parses
    as nothing, and a run that silently reports no decision is worse than one
    that says why.
    """
    from langchain_core.messages import HumanMessage, SystemMessage
    resp = llm.invoke([SystemMessage(content=system), HumanMessage(content=user)])
    text = resp.content if hasattr(resp, "content") else str(resp)
    if isinstance(text, list):        # content blocks → plain text
        text = "".join(b.get("text", "") if isinstance(b, dict) else str(b) for b in text)
    usage = getattr(resp, "usage_metadata", None) or {}
    meta = getattr(resp, "response_metadata", {}) or {}
    if not usage:
        usage = meta.get("usage", {}) or {}
    tin = int(usage.get("input_tokens") or usage.get("prompt_tokens") or 0)
    tout = int(usage.get("output_tokens") or usage.get("completion_tokens") or 0)
    if meta.get("stop_reason") == "max_tokens":
        raise TruncatedReplyError(
            f"the model reached its {tout:,}-token output limit mid-answer")
    return text, tin, tout


_JSON_FENCE = re.compile(r"```(?:json)?\s*(.+?)\s*```", re.S)


def _parse_json(text: str) -> Any | None:
    """Best-effort JSON out of a model reply (fenced, bare, or embedded)."""
    for candidate in (
        (_JSON_FENCE.search(text) or [None, None])[1],
        text.strip(),
        text[text.find("{"): text.rfind("}") + 1] if "{" in text and "}" in text else None,
    ):
        if not candidate:
            continue
        try:
            return json.loads(candidate)
        except (json.JSONDecodeError, TypeError):
            continue
    return None


# ─────────────────────────────────────────────────────────────────────────────
# The two agents
#
# Exactly two agents run a job, both driven by role cards in claude_agents/:
#
#   PLANNER    spawned once, invoked once — turns the loan product, the
#              operator prompt and the document inventory into a processing plan
#              (it never sees document contents).
#   PROCESSOR  spawned once, invoked N+1 times — once per document following the
#              plan, then once more to decide eligibility across the whole file.
#
# Unlike the pipeline's _EphemeralAgent (single-use), the processor is a durable
# identity across its calls, so one agent owns the whole document set and its
# token spend is attributed in one place.
# ─────────────────────────────────────────────────────────────────────────────

AGENT_ROLES: dict[str, dict[str, dict[str, Any]]] = {
    "loan": {
        "planner":   {"agent_id": "LOAN_PLANNER",   "label": "Planning Agent",
                      "card": "CLAUDE_LOAN_PLANNER.md",   "max_tokens": 2048},
        "processor": {"agent_id": "LOAN_PROCESSOR", "label": "Processing Agent",
                      # 4096 truncated the assessment on a three-applicant file
                      # once policy clauses were cited; the cap is a ceiling,
                      # not a charge, so raising it costs nothing unused.
                      "card": "CLAUDE_LOAN_PROCESSOR.md", "max_tokens": 8192},
    },
    "account": {
        "planner":   {"agent_id": "ACCOUNT_PLANNER",   "label": "Planning Agent",
                      "card": "CLAUDE_ACCOUNT_PLANNER.md",   "max_tokens": 2048},
        "processor": {"agent_id": "ACCOUNT_PROCESSOR", "label": "Processing Agent",
                      "card": "CLAUDE_ACCOUNT_PROCESSOR.md", "max_tokens": 8192},
    },
}

MAX_AGENTS_PER_JOB = 2

# ── Cost model ───────────────────────────────────────────────────────────────
# USD per MILLION tokens, (input, output), per Anthropic list pricing.
# A model that isn't listed yields cost=None — the UI shows "—" rather than a
# number we can't stand behind. Local providers (ollama) are free.
PRICING_USD_PER_MTOK: dict[str, tuple[float, float]] = {
    "claude-haiku-4-5":            (1.00,  5.00),
    "claude-haiku-4-5-20251001":   (1.00,  5.00),
    "claude-sonnet-4-6":           (3.00, 15.00),
    "claude-sonnet-5":             (3.00, 15.00),
    "claude-opus-4-8":             (5.00, 25.00),
    "claude-opus-5":               (5.00, 25.00),
    "claude-fable-5":             (10.00, 50.00),
}


def cost_usd(tokens_in: int, tokens_out: int,
             model: str = WORKER_MODEL) -> float | None:
    """Cost of a token spend at list price, or None if the model has no rate."""
    if LLM_PROVIDER == "ollama":
        return 0.0
    rates = PRICING_USD_PER_MTOK.get(model)
    if not rates:
        return None
    rate_in, rate_out = rates
    return (tokens_in / 1_000_000) * rate_in + (tokens_out / 1_000_000) * rate_out


def _load_card(filename: str) -> str:
    path = AGENTS_DIR / filename
    if not path.exists():
        raise LoanConfigError(f"Agent role card not found: {path}")
    return path.read_text(encoding="utf-8")


class LoanAgent:
    """A named, durable agent: one role card, one model client, many calls."""

    def __init__(self, role: str, job: "LoanJob") -> None:
        spec = AGENT_ROLES.get(job.domain, AGENT_ROLES["loan"])[role]
        self.role = role
        self.agent_id: str = spec["agent_id"]
        self.label: str = spec["label"]
        self.card: str = spec["card"]
        self.system: str = _load_card(spec["card"])
        self.model: str = WORKER_MODEL
        self._llm = _make_llm(spec["max_tokens"])
        self.status = "alive"
        self.calls = 0
        self.tokens_in = 0
        self.tokens_out = 0
        self._job = job

    def invoke(self, user_message: str) -> str:
        if self.status != "alive":
            raise RuntimeError(f"[{self.agent_id}] already torn down")
        text, tin, tout = _invoke(self._llm, self.system, user_message)
        self.calls += 1
        self.tokens_in += tin
        self.tokens_out += tout
        self._job.tokens_in += tin
        self._job.tokens_out += tout
        self._job.emit("agent_state", agent=self.public(),
                       tokens_in=self._job.tokens_in,
                       tokens_out=self._job.tokens_out,
                       cost_usd=self._job.cost,
                       elapsed_s=self._job.elapsed_s)
        return text

    def teardown(self) -> None:
        self._llm = None                                   # type: ignore[assignment]
        self.system = ""
        self.status = "torn_down"
        self._job.emit("agent_state", agent=self.public(),
                       tokens_in=self._job.tokens_in,
                       tokens_out=self._job.tokens_out)
        log.info("[%s] %s torn down — %d calls, %d in / %d out tokens",
                 self._job.job_id, self.agent_id, self.calls,
                 self.tokens_in, self.tokens_out)

    def public(self) -> dict:
        return {"agent_id": self.agent_id, "role": self.role, "label": self.label,
                "status": self.status, "calls": self.calls, "model": self.model,
                "tokens_in": self.tokens_in, "tokens_out": self.tokens_out,
                "cost_usd": cost_usd(self.tokens_in, self.tokens_out, self.model),
                "card": self.card}


class AgentTeam:
    """Spawns and tracks a job's agents, capped at MAX_AGENTS_PER_JOB."""

    def __init__(self, job: "LoanJob") -> None:
        self._job = job
        self.agents: list[LoanAgent] = []

    def spawn(self, role: str) -> LoanAgent:
        if len(self.agents) >= MAX_AGENTS_PER_JOB:
            raise LoanConfigError(
                f"Agent budget exhausted ({MAX_AGENTS_PER_JOB} per job): "
                f"cannot spawn '{role}'."
            )
        agent = LoanAgent(role, self._job)
        self.agents.append(agent)
        self._job.agent_records = self.agents
        log.info("[%s] spawned %s (%s) with %s  [%d/%d]", self._job.job_id,
                 agent.agent_id, agent.card, agent.model,
                 len(self.agents), MAX_AGENTS_PER_JOB)
        self._job.emit("agent_spawned", agent=agent.public())
        return agent

    def teardown_all(self) -> None:
        for a in self.agents:
            if a.status == "alive":
                a.teardown()


# ── prompts handed to the agents (the role cards carry the output contracts) ──

def _planner_facts_message(job: "LoanJob") -> str:
    """Deterministic mode: the planner reasons over facts, not raw documents."""
    return (f"LOAN PRODUCT: {job.loan_label}\n\n"
            f"OPERATOR PROMPT (carries the eligibility criteria):\n{job.prompt}\n\n"
            f"DETERMINISTIC EXTRACTION ALREADY COMPLETED — every figure below was "
            f"parsed and arithmetically reconciled in code, not inferred:\n"
            f"{json.dumps(job.fact_sheet, indent=2, default=str)[:MAX_PASS2_CHARS]}\n\n"
            f"Plan what still needs judgement. Code has already settled the "
            f"arithmetic; do not re-derive it.")


def _exception_message(job: "LoanJob", entry: dict) -> str:
    """Row-scoped when the parser understood the document; full text when it
    didn't. An unrecognised layout must not reach the model as a bare error."""
    if entry.get("full_text"):
        return (f"MODE: DOCUMENT\n\n"
                f"OPERATOR PROMPT:\n{job.prompt}\n\n"
                f"No deterministic parser matched this document's layout, so it "
                f"is given to you in full.\n\n"
                f"--- DOCUMENT: {entry['document']} ---\n{entry['full_text']}")
    return (f"MODE: DOCUMENT\n\n"
            f"OPERATOR PROMPT:\n{job.prompt}\n\n"
            f"A deterministic parser read this document and reconciled every "
            f"figure except the rows below. Explain what went wrong on these "
            f"rows and whether the document can be trusted. You are NOT being "
            f"given the whole document — only its failures.\n\n"
            f"--- DOCUMENT: {entry['document']} ({entry['type']}) ---\n"
            f"{json.dumps(entry['failures'], indent=2, default=str)}")


def _facts_assessment_message(job: "LoanJob", resolutions: list[dict]) -> str:
    bundle, _ = _truncate(json.dumps(job.fact_sheet, indent=2, default=str),
                          MAX_PASS2_CHARS)
    return (f"MODE: ASSESSMENT\n\n"
            f"OPERATOR PROMPT:\n{job.prompt}{job.policy_context}\n\n"
            f"PLAN:\n{json.dumps(job.plan, indent=2) if job.plan else '(none)'}\n\n"
            f"DETERMINISTIC FACTS (parsed and reconciled in code — treat every "
            f"figure and every criterion marked 'deterministic' as settled; do "
            f"not recompute them):\n{bundle}\n\n"
            + (f"EXCEPTION FINDINGS (your own analysis of the failing rows):\n"
               f"{json.dumps(resolutions, indent=2, default=str)}\n\n"
               if resolutions else "")
            + "Return the eligibility decision. Carry the code-computed criteria "
              "through verbatim, add any criterion only judgement can settle, and "
              "explain the decision.")


def _planner_message(job: "LoanJob") -> str:
    inventory = "\n".join(
        f"- {d.name} ({Path(d.path).suffix.lower() or 'no extension'}, "
        f"{_kb(d.path)})" for d in job.docs
    )
    return (f"LOAN PRODUCT: {job.loan_label}\n\n"
            f"OPERATOR PROMPT (carries the eligibility criteria):\n{job.prompt}\n\n"
            f"DOCUMENT INVENTORY ({len(job.docs)} files, names only):\n{inventory}")


def _document_message(job: "LoanJob", doc: "DocResult", body: str,
                      truncated: bool) -> str:
    entry = job.plan_for(doc.name)
    return (f"MODE: DOCUMENT\n\n"
            f"OPERATOR PROMPT:\n{job.prompt}\n\n"
            f"PLAN FOR THIS DOCUMENT:\n{json.dumps(entry, indent=2)}\n\n"
            f"--- DOCUMENT: {doc.name} ---\n{body}"
            + ("\n\n[document truncated for length]" if truncated else ""))


def _assessment_message(job: "LoanJob", findings: list[dict],
                        unreadable: list[dict]) -> str:
    bundle, _ = _truncate(json.dumps(findings, indent=2, default=str),
                          MAX_PASS2_CHARS)
    plan = json.dumps(job.plan, indent=2) if job.plan else "(no plan available)"
    return (f"MODE: ASSESSMENT\n\n"
            f"OPERATOR PROMPT:\n{job.prompt}{job.policy_context}\n\n"
            f"PLAN:\n{plan}\n\n"
            f"PER-DOCUMENT FINDINGS ({len(findings)} documents):\n{bundle}\n"
            + (f"\nDOCUMENTS THAT COULD NOT BE READ:\n"
               f"{json.dumps(unreadable, indent=2)}\n" if unreadable else ""))


def _kb(path: str) -> str:
    try:
        return f"{Path(path).stat().st_size / 1024:.1f} KB"
    except OSError:
        return "size unknown"


# ─────────────────────────────────────────────────────────────────────────────
# Job model
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class DocResult:
    name: str
    path: str
    status: str = "pending"        # pending | running | done | failed | skipped
    error: str = ""
    tokens_in: int = 0
    tokens_out: int = 0
    chars: int = 0
    truncated: bool = False
    analysis: Any | None = None
    analysis_text: str = ""

    def public(self) -> dict:
        return {"name": self.name, "status": self.status, "error": self.error,
                "tokens_in": self.tokens_in, "tokens_out": self.tokens_out}


@dataclass
class LoanJob:
    job_id: str
    loan_type: str
    loan_label: str
    input_path: str
    output_path: str          # the parent the operator chose
    prompt: str
    run_dir: str = ""         # the date-stamped folder this run writes into
    bank_name: str = ""       # institution named in the report header/footer
    domain: str = "loan"      # loan | account — drives report wording
    status: str = "queued"          # queued | running | completed | failed | cancelled
    error: str = ""
    total: int = 0
    done: int = 0
    failed: int = 0
    tokens_in: int = 0
    tokens_out: int = 0
    # scanning | planning | documents | assessment | finished
    phase: str = "scanning"
    started_at: str = ""
    finished_at: str = ""
    decision: str = ""
    report_path: str = ""
    plan: dict | None = None
    # deterministic | ai_first — see _execute_deterministic / _execute_ai_first
    mode: str = "deterministic"
    # Optional credit-policy retrieval (loan_policy.py). Empty unless the
    # operator configured a policy pack for this box; policy_context is
    # appended to the assessment prompt only.
    policy_path: str = ""
    policy_context: str = ""
    policy_citations: list[dict] = field(default_factory=list)
    policy_record: dict | None = None
    fact_sheet: dict | None = None
    docs_clean: int = 0            # reconciled in code, zero tokens
    docs_escalated: int = 0        # sent to the Processing Agent
    docs: list[DocResult] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)
    agent_records: list["LoanAgent"] = field(default_factory=list)
    q: queue.Queue = field(default_factory=queue.Queue)
    cancel_event: threading.Event = field(default_factory=threading.Event)
    # Wall-clock timing: _t0 is a monotonic reference, frozen into elapsed_s
    # when the job finishes so a completed card stops ticking.
    _t0: float = field(default_factory=time.perf_counter)
    _elapsed_final: float | None = None

    @property
    def elapsed_s(self) -> float:
        if self._elapsed_final is not None:
            return self._elapsed_final
        return round(time.perf_counter() - self._t0, 1)

    @property
    def cost(self) -> float | None:
        """Total list-price cost of this job so far, summed per agent model."""
        if not self.agent_records:
            return cost_usd(self.tokens_in, self.tokens_out)
        per_agent = [cost_usd(a.tokens_in, a.tokens_out, a.model)
                     for a in self.agent_records]
        return None if any(c is None for c in per_agent) else sum(per_agent)

    def plan_for(self, doc_name: str) -> dict:
        """The planner's entry for one document (empty dict if it has none)."""
        for entry in (self.plan or {}).get("documents", []) or []:
            if isinstance(entry, dict) and entry.get("name") == doc_name:
                return entry
        return {}

    # ── event plumbing ──────────────────────────────────────────────────────
    def emit(self, type_: str, **data) -> None:
        data["type"] = type_
        data["ts"] = datetime.now().isoformat()
        self.q.put(data)

    @property
    def is_cancelled(self) -> bool:
        return self.cancel_event.is_set()

    def snapshot(self) -> dict:
        """Full state — used by /loan/jobs/<id> and by every terminal event."""
        return {
            "job_id": self.job_id, "loan_type": self.loan_type,
            "loan_label": self.loan_label, "status": self.status,
            "phase": self.phase, "error": self.error,
            "total": self.total, "done": self.done, "failed": self.failed,
            "tokens_in": self.tokens_in, "tokens_out": self.tokens_out,
            "cost_usd": self.cost, "elapsed_s": self.elapsed_s,
            "model": WORKER_MODEL, "provider": LLM_PROVIDER,
            "input_path": self.input_path, "output_path": self.output_path,
            "run_dir": self.run_dir, "run_folder": Path(self.run_dir).name
            if self.run_dir else "", "bank_name": self.bank_name,
            "decision": self.decision, "report_path": self.report_path,
            "started_at": self.started_at, "finished_at": self.finished_at,
            "docs": [d.public() for d in self.docs],
            "skipped": self.skipped,
            "agents": [a.public() for a in self.agent_records],
            "plan": self.plan,
            "mode": self.mode,
            # In ai_first mode every document goes to the model by definition,
            # so the counters are derived rather than tallied.
            "docs_clean": self.docs_clean if self.mode == "deterministic" else 0,
            "docs_escalated": (self.docs_escalated if self.mode == "deterministic"
                               else self.done),
            "ai_share": self.ai_share,
            "policy_path": self.policy_path,
            "policy_citations": self.policy_citations,
        }

    @property
    def ai_share(self) -> float:
        """Fraction of documents that required a model call."""
        if not self.total:
            return 0.0
        if self.mode != "deterministic":
            return 1.0
        return round(self.docs_escalated / self.total, 4)


_jobs: dict[str, LoanJob] = {}
_jobs_lock = threading.Lock()
MAX_RETAINED_JOBS = 50


def get_job(job_id: str) -> LoanJob | None:
    with _jobs_lock:
        return _jobs.get(job_id)


def _register(job: LoanJob) -> None:
    with _jobs_lock:
        _jobs[job.job_id] = job
        if len(_jobs) > MAX_RETAINED_JOBS:
            for jid, j in sorted(_jobs.items(), key=lambda kv: kv[1].started_at):
                if j.status in ("completed", "failed", "cancelled"):
                    del _jobs[jid]
                if len(_jobs) <= MAX_RETAINED_JOBS:
                    break


# ─────────────────────────────────────────────────────────────────────────────
# Job execution
# ─────────────────────────────────────────────────────────────────────────────

def _make_run_dir(parent: Path, loan_type: str) -> Path:
    """A date-stamped folder per run, so reports never overwrite each other.

    Same shape as the orchestrator's project dirs: <YYYYMMDD_HHMMSS>_<slug>.
    Two runs started in the same second would collide, so the name is claimed
    atomically with mkdir(exist_ok=False) and suffixed on collision.
    """
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base = f"{stamp}_{loan_type}"
    candidate = parent / base
    suffix = 2
    while True:
        try:
            candidate.mkdir(parents=True, exist_ok=False)
            return candidate
        except FileExistsError:
            candidate = parent / f"{base}_{suffix}"
            suffix += 1


def start_job(loan_type: str, input_path: str, output_path: str,
              prompt: str, mode: str = "deterministic",
              bank_name: str = "", policy_path: str = "") -> LoanJob:
    """Validate inputs, register the job, and run it on a daemon thread.

    Raises LoanConfigError for anything the operator can fix in the form.
    """
    spec = type_spec(loan_type)
    if not spec:
        raise LoanConfigError(f"Unknown processing type: {loan_type!r}")
    label = spec["label"]

    src = Path(input_path).expanduser()
    if not input_path.strip():
        raise LoanConfigError("Input path is required.")
    if not src.exists():
        raise LoanConfigError(f"Input path not found: {src}")

    if not output_path.strip():
        raise LoanConfigError("Output path is required.")
    out = Path(output_path).expanduser()
    if out.exists() and not out.is_dir():
        raise LoanConfigError(f"Output path is not a directory: {out}")
    try:
        out.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise LoanConfigError(f"Cannot create output path: {exc}") from exc

    prompt = (prompt or "").strip() or DEFAULT_PROMPTS.get(loan_type, "")
    if len(prompt) < 20:
        raise LoanConfigError("Prompt must be at least 20 characters.")

    supported, skipped = scan_documents(src)
    if not supported:
        raise LoanConfigError(
            f"No processable documents in {src} — supported types: "
            + ", ".join(sorted(SUPPORTED))
        )

    if mode not in ("deterministic", "ai_first"):
        raise LoanConfigError("mode must be 'deterministic' or 'ai_first'")

    # A configured pack must exist before the run starts — failing here is a
    # form error the operator can fix, rather than a silent miss mid-run.
    policy_path = (policy_path or "").strip()
    if policy_path and not Path(policy_path).expanduser().exists():
        raise LoanConfigError(f"Policy pack not found: {policy_path}")

    try:
        run_dir = _make_run_dir(out, loan_type)
    except OSError as exc:
        raise LoanConfigError(f"Cannot create the run folder: {exc}") from exc

    job = LoanJob(
        job_id=f"{spec['domain']}-{loan_type}-{uuid.uuid4().hex[:8]}",
        loan_type=loan_type, loan_label=label, mode=mode,
        domain=spec["domain"],
        bank_name=(bank_name or "").strip(),
        policy_path=policy_path,
        input_path=str(src), output_path=str(out), run_dir=str(run_dir),
        prompt=prompt,
        total=len(supported), started_at=datetime.now().isoformat(),
        docs=[DocResult(name=str(p.relative_to(src) if src.is_dir() else p.name),
                        path=str(p)) for p in supported],
        skipped=[p.name for p in skipped[:20]],
    )
    _register(job)

    threading.Thread(target=_run_job, args=(job,), daemon=True,
                     name=f"loanjob-{job.job_id}").start()
    return job


def _run_job(job: LoanJob) -> None:
    job.status = "running"
    job.emit("job_started", **job.snapshot())
    t0 = time.perf_counter()
    team = AgentTeam(job)

    try:
        _execute(job, team, t0)
    except LoanConfigError as exc:
        _finish(job, "failed", str(exc))
    except Exception as exc:                                          # noqa: BLE001
        log.exception("[%s] job crashed", job.job_id)
        _finish(job, "failed", f"{type(exc).__name__}: {exc}")
    finally:
        team.teardown_all()


def _execute(job: LoanJob, team: AgentTeam, t0: float) -> None:
    if job.mode == "deterministic":
        _execute_deterministic(job, team, t0)
    else:
        _execute_ai_first(job, team, t0)


def _attach_policy(job: LoanJob) -> None:
    """Retrieve the governing clauses for this job, once, before the assessment.

    A no-op unless the box has a policy pack configured. Best-effort: a pack
    that is missing, unreadable, or returns nothing leaves the assessment
    exactly as it would have been without the feature.
    """
    if not job.policy_path:
        return
    try:
        import loan_policy
    except Exception as exc:                                      # noqa: BLE001
        log.warning("[%s] policy retrieval unavailable: %s", job.job_id, exc)
        return

    # Code first: where the policy states a number and the extractors parsed
    # the matching figure, the band is selected and the ratio compared here —
    # no tokens, no model arithmetic. Retrieval is then narrowed to whatever is
    # genuinely left open (documentation, tenure, conduct).
    coded = loan_policy.apply_in_code(job.policy_path, job.fact_sheet or {})

    ctx = loan_policy.retrieve(
        job.policy_path, job.loan_label, job.prompt,
        k=loan_policy.RESIDUAL_K if coded.ok else loan_policy.DEFAULT_K,
        agent_id=f"{job.domain.upper()}_PROCESSOR",
        exclude=loan_policy.SETTLED_TERMS if coded.ok else (),
        exclude_sources=coded.sources_settled,
    )

    job.policy_context = coded.context + ctx.context
    job.policy_citations = ctx.citations
    job.policy_record = {
        **ctx.record(),
        "code_side": {
            "rules_origin": coded.rules_origin,
            "rules_found": coded.rules_found,
            "applicants": coded.applicants,
            "elapsed_ms": round(coded.elapsed_ms, 1),
            "checks": coded.checks,
        },
    }

    # Code-settled policy checks are criteria like any other the pipeline
    # computed, so they join the fact sheet and flow into the report's table.
    if coded.ok and isinstance(job.fact_sheet, dict):
        job.fact_sheet.setdefault("criteria_computed_in_code", []).extend(coded.lean_checks)

    if job.run_dir:
        try:
            (Path(job.run_dir) / "policy_retrieval.json").write_text(
                json.dumps(job.policy_record, indent=2), encoding="utf-8")
        except OSError as exc:
            log.warning("[%s] could not write policy_retrieval.json: %s",
                        job.job_id, exc)

    job.emit("policy_retrieved", clauses=len(ctx.citations),
             citations=ctx.citations, retrieve_ms=round(ctx.retrieve_ms, 1),
             chunks_in_pack=ctx.chunks_in_pack,
             coded_checks=len(coded.checks),
             coded_settled=sum(1 for c in coded.checks
                               if c["status"] in ("met", "not_met")))


# ─────────────────────────────────────────────────────────────────────────────
# Deterministic-first path — parse and reconcile in code, escalate only failures
# ─────────────────────────────────────────────────────────────────────────────

def _execute_deterministic(job: LoanJob, team: AgentTeam, t0: float) -> None:
    import loan_extractors as X

    # ── pass 1: zero-token extraction + arithmetic validation ───────────────
    job.phase = "extracting"
    job.emit("phase_changed", phase="extracting")
    results: list[X.ExtractionResult] = []

    for idx, doc in enumerate(job.docs):
        if job.is_cancelled:
            _finish(job, "cancelled", "Cancelled by the operator.")
            return
        job.emit("doc_started", index=idx, name=doc.name,
                 done=job.done, total=job.total)

        res = X.extract_document(Path(doc.path))
        results.append(res)
        doc.analysis = res.public()
        job.done += 1
        if res.status == "clean":
            doc.status = "done"
            job.docs_clean += 1
        else:
            doc.status = "escalated"
            reasons = [e.get("reason", "") for e in res.exceptions[:2]]
            doc.error = ("no parser for this layout — sent to the model in full"
                         if reasons == ["no_parser_for_layout"]
                         else "; ".join(r.replace("_", " ") for r in reasons))
            job.docs_escalated += 1
        job.emit("doc_completed", index=idx, **_doc_event(job, doc))

    job.fact_sheet = X.fact_sheet(results,
                                  include_cross_checks=job.domain == "loan")

    # A document no parser recognised still needs its text in front of the
    # model — otherwise deterministic mode would silently downgrade unknown
    # layouts to "here is an error message".
    by_name = {d.name: d for d in job.docs}
    for entry in job.fact_sheet.get("exceptions", []):
        res = next((r for r in results if r.document == entry["document"]), None)
        if res is None or res.status != "unreadable":
            continue
        doc = by_name.get(entry["document"]) or next(
            (d for d in job.docs if Path(d.path).name == entry["document"]), None)
        if not doc:
            continue
        text, err = extract_text(Path(doc.path))
        if text:
            body, truncated = _truncate(text, MAX_DOC_CHARS)
            entry["full_text"] = body + (
                "\n\n[document truncated for length]" if truncated else "")
        else:
            entry["failures"].append({"reason": "text_extraction_failed",
                                      "error": err})

    _write_json(Path(job.run_dir) / "extraction_facts.json", job.fact_sheet)
    docs_dir = Path(job.run_dir) / "documents"
    docs_dir.mkdir(parents=True, exist_ok=True)
    for res in results:
        _write_json(docs_dir / f"{Path(res.document).stem}.json", res.public())

    log.info("[%s] deterministic pass: %d/%d clean, %d escalated (0 tokens)",
             job.job_id, job.docs_clean, job.total, job.docs_escalated)

    if job.is_cancelled:
        _finish(job, "cancelled", "Cancelled by the operator.")
        return

    # ── agent 1 of 2: the planner, reasoning over facts rather than documents ─
    job.phase = "planning"
    job.emit("phase_changed", phase="planning")
    planner = team.spawn("planner")
    try:
        reply = planner.invoke(_planner_facts_message(job))
        parsed = _parse_json(reply)
        job.plan = parsed if isinstance(parsed, dict) else None
        if job.plan is None:
            _write_text(Path(job.run_dir) / "plan_raw.txt", reply)
    except LoanConfigError:
        raise
    except Exception as exc:                                      # noqa: BLE001
        log.warning("[%s] planning failed: %s", job.job_id, exc)
        job.error = f"Planning agent failed ({exc}) — assessed without a plan."
    finally:
        planner.teardown()

    if job.plan:
        _write_json(Path(job.run_dir) / "processing_plan.json", job.plan)
    job.emit("plan_ready", plan=job.plan, error=job.error)

    # ── agent 2 of 2: the processor, on exceptions only ─────────────────────
    processor = team.spawn("processor")
    resolutions: list[dict] = []
    exceptions = (job.fact_sheet or {}).get("exceptions", [])

    if exceptions:
        job.phase = "exceptions"
        job.emit("phase_changed", phase="exceptions")
        for entry in exceptions:
            if job.is_cancelled:
                _finish(job, "cancelled", "Cancelled by the operator.")
                return
            job.emit("doc_started", index=-1, name=entry["document"],
                     done=job.done, total=job.total)
            try:
                reply = processor.invoke(_exception_message(job, entry))
                resolutions.append({"document": entry["document"],
                                    "findings": _parse_json(reply) or reply})
            except Exception as exc:                              # noqa: BLE001
                log.warning("[%s] exception resolution failed for %s: %s",
                            job.job_id, entry["document"], exc)
                resolutions.append({"document": entry["document"],
                                    "error": f"{type(exc).__name__}: {exc}"})

    # ── same agent, final call: decision from the computed facts ────────────
    job.phase = "assessment"
    job.emit("eligibility_started", done=job.done, total=job.total)
    _attach_policy(job)
    decision_obj: Any | None = None
    decision_text = ""
    try:
        decision_text = processor.invoke(_facts_assessment_message(job, resolutions))
        decision_obj = _parse_json(decision_text)
        job.decision = _verdict_of(decision_obj)
    except Exception as exc:                                      # noqa: BLE001
        log.warning("[%s] assessment failed: %s", job.job_id, exc)
        job.error = f"Assessment failed: {exc}"
    processor.teardown()

    unreadable = [{"document": d.name, "error": d.error}
                  for d in job.docs if d.status == "failed"]
    report = {
        "job_id": job.job_id, "loan_type": job.loan_type,
        "loan_label": job.loan_label, "input_path": job.input_path,
        "bank_name": job.bank_name, "run_dir": job.run_dir,
        "mode": job.mode,
        "documents_processed": job.done,
        "documents_reconciled_in_code": job.docs_clean,
        "documents_escalated_to_ai": job.docs_escalated,
        "ai_share": round(job.docs_escalated / job.total, 4) if job.total else 0.0,
        "documents_failed": job.failed,
        "unreadable": unreadable,
        "plan": job.plan,
        "agents": [a.public() for a in job.agent_records],
        "deterministic_facts": job.fact_sheet,
        "exception_findings": resolutions,
        "assessment": decision_obj,
        "assessment_text": None if decision_obj else (decision_text or None),
        "tokens": {"input": job.tokens_in, "output": job.tokens_out},
        "cost_usd": job.cost,
        "model": WORKER_MODEL,
        "provider": LLM_PROVIDER,
        "elapsed_s": job.elapsed_s,
        "generated_at": datetime.now().isoformat(),
    }
    _write_reports(job, report)
    _finish(job, "completed", job.error)


def _write_reports(job: LoanJob, report: dict) -> None:
    path = Path(job.run_dir) / "eligibility_report.json"
    _write_json(path, report)
    _write_text(Path(job.run_dir) / "eligibility_report.md",
                _render_markdown(job, report))
    _write_text(Path(job.run_dir) / "eligibility_report.html",
                _render_html(job, report))
    job.report_path = str(path)
    _write_json(Path(job.run_dir) / "summary.json", {
        **job.snapshot(), "prompt": job.prompt,
        "elapsed_s": report["elapsed_s"],
    })


# ─────────────────────────────────────────────────────────────────────────────
# AI-first path — every document goes to the model (the original behaviour)
# ─────────────────────────────────────────────────────────────────────────────

def _execute_ai_first(job: LoanJob, team: AgentTeam, t0: float) -> None:
    # ── agent 1 of 2: the planner ───────────────────────────────────────────
    job.phase = "planning"
    job.emit("phase_changed", phase="planning")
    planner = team.spawn("planner")
    try:
        plan_reply = planner.invoke(_planner_message(job))
        parsed = _parse_json(plan_reply)
        job.plan = parsed if isinstance(parsed, dict) else None
        if job.plan is None:
            _write_text(Path(job.run_dir) / "plan_raw.txt", plan_reply)
            log.warning("[%s] planner output was not JSON — saved as plan_raw.txt",
                        job.job_id)
    except LoanConfigError:
        raise
    except Exception as exc:                                          # noqa: BLE001
        # A failed plan degrades the run, it does not end it: the processor can
        # still work from the operator prompt alone.
        log.warning("[%s] planning failed: %s", job.job_id, exc)
        job.error = f"Planning agent failed ({exc}) — processed without a plan."
    finally:
        planner.teardown()

    if job.plan:
        _write_json(Path(job.run_dir) / "processing_plan.json", job.plan)
    job.emit("plan_ready", plan=job.plan, error=job.error)

    if job.is_cancelled:
        _finish(job, "cancelled", "Cancelled by the operator.")
        return

    # ── agent 2 of 2: the processor ─────────────────────────────────────────
    job.phase = "documents"
    job.emit("phase_changed", phase="documents")
    processor = team.spawn("processor")

    docs_dir = Path(job.run_dir) / "documents"
    docs_dir.mkdir(parents=True, exist_ok=True)

    for idx, doc in enumerate(job.docs):
        if job.is_cancelled:
            _finish(job, "cancelled", "Cancelled by the operator.")
            return

        doc.status = "running"
        job.emit("doc_started", index=idx, name=doc.name,
                 done=job.done, total=job.total)

        text, err = extract_text(Path(doc.path))
        if err:
            doc.status, doc.error = "failed", err
            job.failed += 1
            job.done += 1
            job.emit("doc_completed", index=idx, **_doc_event(job, doc))
            continue

        body, truncated = _truncate(text, MAX_DOC_CHARS)
        doc.chars, doc.truncated = len(text), truncated

        before_in, before_out = processor.tokens_in, processor.tokens_out
        try:
            reply = processor.invoke(_document_message(job, doc, body, truncated))
        except Exception as exc:                                      # noqa: BLE001
            doc.status, doc.error = "failed", f"{type(exc).__name__}: {exc}"
            job.failed += 1
            job.done += 1
            log.warning("[%s] %s failed: %s", job.job_id, doc.name, exc)
            job.emit("doc_completed", index=idx, **_doc_event(job, doc))
            continue

        doc.tokens_in = processor.tokens_in - before_in
        doc.tokens_out = processor.tokens_out - before_out
        doc.analysis = _parse_json(reply)
        doc.analysis_text = reply
        doc.status = "done"
        job.done += 1

        _write_json(docs_dir / f"{Path(doc.name).stem}.json", {
            "document": doc.name, "source_path": doc.path,
            "loan_type": job.loan_type, "chars": doc.chars,
            "truncated": doc.truncated,
            "processed_by": processor.agent_id,
            "plan_entry": job.plan_for(doc.name) or None,
            "tokens": {"input": doc.tokens_in, "output": doc.tokens_out},
            "analysis": doc.analysis,
            "analysis_text": None if doc.analysis else reply,
            "processed_at": datetime.now().isoformat(),
        })
        job.emit("doc_completed", index=idx, **_doc_event(job, doc))

    if job.is_cancelled:
        _finish(job, "cancelled", "Cancelled by the operator.")
        return

    # ── same agent, final call: eligibility across the whole file ───────────
    job.phase = "assessment"
    job.emit("eligibility_started", done=job.done, total=job.total)
    _attach_policy(job)

    findings = [
        {"document": d.name,
         "findings": d.analysis if d.analysis is not None else d.analysis_text}
        for d in job.docs if d.status == "done"
    ]
    unreadable = [{"document": d.name, "error": d.error}
                  for d in job.docs if d.status == "failed"]

    decision_obj: Any | None = None
    decision_text = ""
    if findings:
        try:
            decision_text = processor.invoke(
                _assessment_message(job, findings, unreadable))
            decision_obj = _parse_json(decision_text)
            job.decision = _verdict_of(decision_obj)
        except Exception as exc:                                      # noqa: BLE001
            log.warning("[%s] assessment failed: %s", job.job_id, exc)
            job.error = f"Assessment failed: {exc}"
    else:
        job.error = "No document could be read — no eligibility assessment was made."

    processor.teardown()

    report = {
        "job_id": job.job_id, "loan_type": job.loan_type,
        "loan_label": job.loan_label, "input_path": job.input_path,
        "bank_name": job.bank_name, "run_dir": job.run_dir,
        "mode": job.mode,
        "documents_processed": job.done - job.failed,
        "documents_reconciled_in_code": 0,
        "documents_escalated_to_ai": job.done,
        "ai_share": 1.0,
        "documents_failed": job.failed,
        "unreadable": unreadable,
        "plan": job.plan,
        "agents": [a.public() for a in job.agent_records],
        "assessment": decision_obj,
        "assessment_text": None if decision_obj else (decision_text or None),
        "tokens": {"input": job.tokens_in, "output": job.tokens_out},
        "cost_usd": job.cost,
        "model": WORKER_MODEL,
        "provider": LLM_PROVIDER,
        "elapsed_s": job.elapsed_s,
        "generated_at": datetime.now().isoformat(),
    }
    report_path = Path(job.run_dir) / "eligibility_report.json"
    _write_json(report_path, report)
    _write_text(Path(job.run_dir) / "eligibility_report.md",
                _render_markdown(job, report))
    _write_text(Path(job.run_dir) / "eligibility_report.html",
                _render_html(job, report))
    job.report_path = str(report_path)

    _write_json(Path(job.run_dir) / "summary.json", {
        **job.snapshot(),
        "prompt": job.prompt,
        "elapsed_s": report["elapsed_s"],
    })

    _finish(job, "completed", job.error)


def _doc_event(job: LoanJob, doc: DocResult) -> dict:
    return {"name": doc.name, "status": doc.status, "error": doc.error,
            "done": job.done, "total": job.total, "failed": job.failed,
            "tokens_in": job.tokens_in, "tokens_out": job.tokens_out,
            "cost_usd": job.cost, "elapsed_s": job.elapsed_s}


def _finish(job: LoanJob, status: str, error: str = "") -> None:
    job.status = status
    job.error = error
    job.phase = "finished"
    job._elapsed_final = round(time.perf_counter() - job._t0, 1)   # stop the clock
    job.finished_at = datetime.now().isoformat()
    event = {"completed": "job_completed", "cancelled": "job_cancelled"}.get(
        status, "job_failed")
    job.emit(event, **job.snapshot())
    job.emit("stream_end")
    log.info("[%s] %s (%d/%d docs, %d in / %d out tokens)", job.job_id, status,
             job.done - job.failed, job.total, job.tokens_in, job.tokens_out)


# ─────────────────────────────────────────────────────────────────────────────
# Output writers
# ─────────────────────────────────────────────────────────────────────────────

def fmt_cost(value: float | None) -> str:
    """Money for humans: small spends need more decimals, not fewer."""
    if value is None:
        return "—"
    if value == 0:
        return "$0.00 (local model)"
    if value >= 1:
        return f"${value:,.2f}"
    if value >= 0.01:
        return f"${value:.3f}"
    return f"${value:.4f}"


def fmt_duration(seconds: float | None) -> str:
    if not seconds:
        return "0s"
    seconds = round(seconds)
    if seconds < 60:
        return f"{seconds}s"
    minutes, secs = divmod(seconds, 60)
    if minutes < 60:
        return f"{minutes}m {secs:02d}s"
    hours, minutes = divmod(minutes, 60)
    return f"{hours}h {minutes:02d}m"


def _write_json(path: Path, payload: Any) -> None:
    try:
        path.write_text(json.dumps(payload, indent=2, default=str), encoding="utf-8")
    except OSError as exc:
        log.warning("could not write %s: %s", path, exc)


def _write_text(path: Path, text: str) -> None:
    try:
        path.write_text(text, encoding="utf-8")
    except OSError as exc:
        log.warning("could not write %s: %s", path, exc)


def _render_markdown(job: LoanJob, report: dict) -> str:
    doc_title, verdict_word = report_titles(job)
    bank = report.get("bank_name") or ""
    lines = [
        f"# {job.loan_label} — {doc_title.lower()}",
        "",
        *([f"**{bank}** · Governed loan document processing", ""] if bank else []),
        f"- **Job**: `{job.job_id}`",
        f"- **Input**: `{job.input_path}`",
        f"- **Documents**: {report['documents_processed']} processed"
        + (f", {report['documents_failed']} unreadable" if report["documents_failed"] else ""),
        f"- **Tokens**: {job.tokens_in:,} in / {job.tokens_out:,} out",
        f"- **Cost**: {fmt_cost(report.get('cost_usd'))} at list price "
        f"({report.get('model', '')})",
        f"- **Time**: {fmt_duration(report.get('elapsed_s', 0))}",
        f"- **Mode**: {report.get('mode', 'ai_first')} — "
        f"{report.get('documents_reconciled_in_code', 0)} reconciled in code, "
        f"{report.get('documents_escalated_to_ai', 0)} escalated to AI "
        f"({report.get('ai_share', 1.0) * 100:.0f}% AI share)",
        f"- **Generated**: {report['generated_at']}",
        "",
    ]
    if not isinstance(report.get("assessment"), dict):
        lines += ["## Assessment", "", report.get("assessment_text")
                  or "_No assessment was produced._", ""]
        return "\n".join(lines)

    a = normalise_assessment(report.get("assessment"))
    subject_label = "Customer" if job.domain == "account" else "Applicant"
    lines += [f"## {verdict_word}: {a['outcome'] or 'UNKNOWN'}",
              "",
              f"**{subject_label}**: {a['subject']}  ",
              f"**Confidence**: {a['confidence'] or 'unknown'}"
              + (f"  \n**Period**: {a['period']}" if a["period"] else ""),
              "",
              a["rationale"], ""]

    if a["findings"]:
        lines += ["## Findings", "", "| Item | Value | Basis |", "|---|---|---|"]
        for f in a["findings"]:
            if isinstance(f, dict):
                lines.append(f"| {f.get('label','')} | {f.get('value','')} "
                             f"| {f.get('basis','')} |")
        lines.append("")

    criteria = a["checks"]
    if criteria:
        heading = "Checks" if job.domain == "account" else "Criteria"
        lines += [f"## {heading}", "", "| Check | Status | Evidence |",
                  "|---|---|---|"]
        for c in criteria:
            if isinstance(c, dict):
                lines.append(f"| {c.get('criterion') or c.get('check','')} "
                             f"| {c.get('status','')} | {c.get('evidence','')} |")
        lines.append("")

    for title, key in (("Missing documents", "missing_documents"),
                       ("Risk flags", "risk_flags"),
                       ("Next steps", "next_steps")):
        items = a.get(key) or []
        if items:
            lines += [f"## {title}", ""] + [f"- {i}" for i in items] + [""]

    if report.get("unreadable"):
        lines += ["## Unreadable documents", ""] + [
            f"- `{u['document']}` — {u['error']}" for u in report["unreadable"]
        ] + [""]

    coded_checks = ((job.policy_record or {}).get("code_side") or {}).get("checks") or []
    if coded_checks:
        lines += ["## Policy thresholds applied in code", "",
                  "Bands selected from the policy and compared against the parsed "
                  "figures by the pipeline — no model arithmetic.", "",
                  "| Criterion | Status | Basis | Clause |", "|---|---|---|---|"]
        for c in coded_checks:
            lines.append(f"| {c['criterion']} | {c['status']} | {c['evidence']} "
                         f"| `{c['policy_source']}` {c['clause']} |")
        lines.append("")

    if job.policy_citations:
        lines += ["## Credit policy retrieved", "",
                  "Clauses retrieved from the bank's policy pack and supplied to "
                  "the Processing Agent for this assessment.", "",
                  "| Source | Clause | Score |", "|---|---|---|"]
        for c in job.policy_citations:
            lines.append(f"| `{c['source']}` | {c['span']} | {c['score']:.2f} |")
        lines.append("")

    lines += ["---", "",
              f"{bank + ' — ' if bank else ''}{job.loan_label} {doc_title.lower()}, "
              f"generated {report.get('generated_at', '')}. For internal credit "
              f"review only.",
              f"Written to `{job.run_dir or job.output_path}`.", ""]
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# HTML report — same document language as the batch processing reports
# (see sample_report_ai_repaired.html): print-friendly, rule-lined, stamped.
# ─────────────────────────────────────────────────────────────────────────────

_STAMP = {
    "ELIGIBLE":     ("#1E6B4E", "ELIGIBLE", "All criteria evidenced by the file"),
    "NOT_ELIGIBLE": ("#A33B2E", "NOT ELIGIBLE", "A criterion is breached on the evidence"),
    "NEEDS_REVIEW": ("#8A5A00", "NEEDS REVIEW", "Criteria unverified — human review required"),
    # Account processing outcomes
    "COMPLETE":     ("#1E6B4E", "COMPLETE", "Every required item is present and valid"),
    "INCOMPLETE":   ("#A33B2E", "INCOMPLETE", "Required items are missing or expired"),
    "REPORTED":     ("#185FA5", "REPORTED", "Figures reported from the documents supplied"),
}


def _verdict_of(assessment: Any) -> str:
    """The one-word answer, whichever domain produced it.

    Loans answer with `decision` (ELIGIBLE / NOT_ELIGIBLE); account work answers
    with `outcome` (COMPLETE / REPORTED / NEEDS_REVIEW …). Both land on the same
    job field so the box, the SSE events and the report agree.
    """
    if not isinstance(assessment, dict):
        return ""
    return (assessment.get("decision") or assessment.get("outcome") or "").strip()


def normalise_assessment(assessment: Any) -> dict:
    """One shape for both domains.

    Loans answer with decision/criteria/applicant; account work answers with
    outcome/checks/subject plus reported findings. The report renders one
    structure, so the two are mapped onto it here rather than in the templates.
    """
    a = assessment if isinstance(assessment, dict) else {}
    return {
        "subject": a.get("applicant") or a.get("subject") or "unknown",
        "outcome": (a.get("decision") or a.get("outcome") or "").upper(),
        "confidence": a.get("confidence") or "",
        "period": a.get("period"),
        "checks": a.get("criteria") or a.get("checks") or [],
        "findings": a.get("findings") or [],
        "missing_documents": a.get("missing_documents") or [],
        "risk_flags": a.get("risk_flags") or [],
        "rationale": a.get("rationale") or "",
        "next_steps": a.get("next_steps") or [],
    }


def report_titles(job: "LoanJob") -> tuple[str, str]:
    """(document title, the word used for the verdict section)."""
    if job.domain == "account":
        return "Processing report", "Outcome"
    return "Eligibility report", "Decision"


def _esc(value: Any) -> str:
    return (str("" if value is None else value)
            .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            .replace('"', "&quot;"))


def _fmt_dt(iso: str) -> str:
    try:
        return datetime.fromisoformat(iso).strftime("%d %b %Y, %H:%M:%S")
    except (ValueError, TypeError):
        return iso or "—"


def _render_html(job: LoanJob, report: dict) -> str:
    a = normalise_assessment(report.get("assessment"))
    doc_title, verdict_word = report_titles(job)
    decision = a["outcome"] or (job.decision or "NEEDS_REVIEW").upper()
    colour, stamp_text, stamp_sub = _STAMP.get(
        decision, ("#66707F", decision or "NO DECISION", "No structured decision was returned"))
    applicant = a["subject"]
    payload_sha = hashlib.sha256(
        json.dumps(report, sort_keys=True, default=str).encode("utf-8")).hexdigest()

    # ── criteria ────────────────────────────────────────────────────────────
    crit_rows = ""
    for c in (a["checks"]):
        if not isinstance(c, dict):
            continue
        status = (c.get("status") or "").lower()
        cls = {"met": "ok", "not_met": "x"}.get(status, "warn")
        mark = {"met": "✓ met", "not_met": "✗ not met"}.get(status, "◦ unverified")
        crit_rows += (f"<tr><td>{_esc(c.get('criterion'))}</td>"
                      f"<td><span class='{cls}'>{mark}</span></td>"
                      f"<td class='narr small'>{_esc(c.get('evidence')) or '—'}</td></tr>")
    crit_block = (f"<h2>{'Checks' if job.domain == 'account' else 'Eligibility criteria'} ({len(a['checks'])})</h2>"
                  f"<table class='grid'><thead><tr><th>Criterion</th><th>Status</th>"
                  f"<th>Evidence</th></tr></thead><tbody>{crit_rows}</tbody></table>"
                  ) if crit_rows else ""

    # ── reported figures (account work answers with these) ──────────────────
    find_rows = ""
    for f in a["findings"]:
        if not isinstance(f, dict):
            continue
        find_rows += (f"<tr><td>{_esc(f.get('label'))}</td>"
                      f"<td class='num'>{_esc(f.get('value'))}</td>"
                      f"<td class='narr small'>{_esc(f.get('basis')) or '—'}</td></tr>")
    findings_block = (f"<h2>Findings ({len(a['findings'])})</h2>"
                      f"<table class='grid'><thead><tr><th>Item</th>"
                      f"<th style='text-align:right'>Value</th><th>Basis</th>"
                      f"</tr></thead><tbody>{find_rows}</tbody></table>"
                      ) if find_rows else ""

    # ── documents ───────────────────────────────────────────────────────────
    doc_rows = ""
    for d in job.docs:
        if d.status == "failed":
            doc_rows += (f"<tr class='rep'><td class='mono small'>{_esc(d.name)}</td>"
                         f"<td colspan='2'><span class='x'>unreadable</span></td>"
                         f"<td class='narr small'>{_esc(d.error)}</td>"
                         f"<td class='num'>—</td></tr>")
            continue
        an = d.analysis if isinstance(d.analysis, dict) else {}
        facts = an.get("key_facts") or []
        facts_html = "<br>".join(_esc(f) for f in facts[:4]) or "—"
        concerns = an.get("concerns") or []
        if concerns:
            facts_html += ("<br><span class='x small'>⚠ "
                           + _esc("; ".join(str(c) for c in concerns[:2])) + "</span>")
        doc_rows += (f"<tr><td class='mono small'>{_esc(d.name)}</td>"
                     f"<td>{_esc(an.get('document_type') or '—')}</td>"
                     f"<td class='small'>{_esc(an.get('relevance') or '—')}</td>"
                     f"<td class='narr small'>{facts_html}</td>"
                     f"<td class='num'>{d.tokens_in + d.tokens_out:,}</td></tr>")

    # ── list sections ───────────────────────────────────────────────────────
    def list_block(title: str, items: list, cls: str = "") -> str:
        rows = "".join(f"<li class='{cls}'>{_esc(i)}</li>" for i in (items or []))
        return f"<h2>{title} ({len(items or [])})</h2><ul class='pts'>{rows}</ul>" if rows else ""

    plan = report.get("plan") or {}
    plan_block = ""
    if plan:
        plan_block = (
            f"<h2>Processing plan · Planning Agent</h2>"
            f"<p class='small'>{_esc(plan.get('applicant_file_summary'))}</p>"
            + list_block("Watch-for signals", plan.get("watch_for") or [])
            + list_block("Expected but absent", plan.get("missing_expected") or [])
        )

    # ── credit policy actually put in front of the agent ────────────────────
    policy_block = ""
    coded_checks = ((job.policy_record or {}).get("code_side") or {}).get("checks") or []
    if coded_checks:
        # Shown from the pipeline's own record, not from the model's answer:
        # these verdicts were computed here, and the report should say so even
        # if the agent paraphrases them.
        rows = "".join(
            f"<tr><td>{_esc(c['criterion'])}</td>"
            f"<td><span class='{'ok' if c['status'] == 'met' else 'x' if c['status'] == 'not_met' else 'muted'}'>"
            f"{_esc(c['status'])}</span></td>"
            f"<td class='narr small'>{_esc(c['evidence'])}</td>"
            f"<td class='mono small'>{_esc(c['policy_source'])} {_esc(c['clause'])}</td></tr>"
            for c in coded_checks)
        decided = sum(1 for c in coded_checks if c["status"] in ("met", "not_met"))
        policy_block += (
            f"<h2>Policy thresholds applied in code ({decided} of "
            f"{len(coded_checks)} decided)</h2>"
            f"<p class='small muted'>Bands selected from the policy and compared "
            f"against the parsed figures by the pipeline — no model arithmetic, "
            f"no tokens. The agent was given these as settled.</p>"
            f"<table class='grid'><thead><tr><th>Criterion</th><th>Status</th>"
            f"<th>Basis</th><th>Clause</th></tr></thead><tbody>{rows}</tbody></table>")

    if job.policy_citations:
        rec = job.policy_record or {}
        pol_rows = "".join(
            f"<tr><td class='mono small'>{_esc(c['source'])}</td>"
            f"<td class='mono small'>{_esc(c['span'])}</td>"
            f"<td class='narr small'>{_esc(c['preview'])}…</td>"
            f"<td class='num'>{c['score']:.2f}</td></tr>"
            for c in job.policy_citations
        )
        policy_block += (
            f"<h2>Credit policy retrieved ({len(job.policy_citations)} clauses)</h2>"
            f"<p class='small muted'>Retrieved from the bank's policy pack and "
            f"supplied to the Processing Agent for this assessment — "
            f"{rec.get('chunks_in_pack', 0):,} clauses in the pack, "
            f"{rec.get('retrieve_ms', 0):.0f} ms, query "
            f"<span class='mono'>{_esc((rec.get('query_sha256') or '')[:12])}</span>.</p>"
            f"<table class='grid'><thead><tr><th>Source</th><th>Clause</th>"
            f"<th>Opening</th><th style='text-align:right'>Score</th></tr></thead>"
            f"<tbody>{pol_rows}</tbody></table>")

    agent_rows = "".join(
        f"<tr><td class='mono small'>{_esc(ag['agent_id'])}</td>"
        f"<td>{_esc(ag['label'])}</td>"
        f"<td class='mono small'>{_esc(ag['card'])}</td>"
        f"<td class='mono small'>{_esc(ag['model'])}</td>"
        f"<td class='num'>{ag['calls']}</td>"
        f"<td class='num'>{ag['tokens_in']:,}</td>"
        f"<td class='num'>{ag['tokens_out']:,}</td>"
        f"<td class='num'>{_esc(fmt_cost(ag.get('cost_usd')))}</td></tr>"
        for ag in (report.get("agents") or [])
    )

    return f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{_esc(doc_title)} — {_esc(job.loan_label)} · {_esc(applicant)}</title>
<style>
:root{{--ink:#1B2431;--rule:#D5DDE9;--faint:#F3F6FA;--muted:#66707F;--stamp:{colour}}}
*{{box-sizing:border-box;margin:0}}
body{{font:14px/1.5 system-ui,-apple-system,"Segoe UI",Roboto,sans-serif;
  color:var(--ink);background:#fff;max-width:920px;margin:0 auto;padding:32px 28px}}
.mono,.num{{font-family:ui-monospace,"Cascadia Mono","SF Mono",Menlo,Consolas,monospace}}
.small{{font-size:12px}}.muted{{color:var(--muted)}}
header{{display:flex;justify-content:space-between;align-items:flex-start;gap:24px;
  border-bottom:3px double var(--ink);padding-bottom:18px;margin-bottom:22px}}
.eyebrow{{font-size:11px;letter-spacing:.18em;text-transform:uppercase;color:var(--muted)}}
h1{{font-size:22px;margin-top:2px}} h1 .mono{{font-size:19px}}
h2{{font-size:12px;letter-spacing:.14em;text-transform:uppercase;
  margin:26px 0 8px;color:var(--muted)}}
.stamp{{border:3px solid var(--stamp);color:var(--stamp);border-radius:8px;
  padding:8px 16px;text-align:center;transform:rotate(-4deg);flex-shrink:0;
  font-weight:800;letter-spacing:.12em;font-size:15px;max-width:230px}}
.stamp small{{display:block;font-weight:400;letter-spacing:0;font-size:10.5px;
  text-transform:none;margin-top:3px}}
.fields{{display:grid;grid-template-columns:repeat(3,1fr);gap:0;border:1px solid var(--rule)}}
.f{{padding:8px 12px;border-right:1px solid var(--rule);border-bottom:1px solid var(--rule)}}
.f:nth-child(3n){{border-right:none}}
.f b{{display:block;font-size:10.5px;letter-spacing:.1em;text-transform:uppercase;
  color:var(--muted);font-weight:600}}
table.grid{{width:100%;border-collapse:collapse;font-size:13px}}
.grid th{{text-align:left;font-size:10.5px;letter-spacing:.1em;text-transform:uppercase;
  color:var(--muted);border-bottom:2px solid var(--ink);padding:6px 8px}}
.grid td{{padding:6px 8px;border-bottom:1px solid var(--rule);vertical-align:top}}
.grid .num{{text-align:right;white-space:nowrap}}
.grid tr.rep td{{background:#FBF4E3}}
.narr{{max-width:360px}}
.ok{{color:#1E6B4E;font-weight:700}}.x{{color:#A33B2E;font-weight:700}}
.warn{{color:#8A5A00;font-weight:700}}
.figs{{display:flex;gap:0;border:1px solid var(--rule);margin-top:6px}}
.figs div{{flex:1;padding:10px 14px;border-right:1px solid var(--rule)}}
.figs div:last-child{{border-right:none}}
.figs b{{display:block;font-size:10.5px;letter-spacing:.1em;text-transform:uppercase;
  color:var(--muted);font-weight:600}}
.figs span{{font-size:16px}}
.rationale{{background:var(--faint);border-left:3px solid var(--stamp);
  padding:12px 14px;margin-top:6px}}
ul.pts{{margin:4px 0 0 18px}} ul.pts li{{padding:2px 0}}
footer{{margin-top:30px;border-top:1px dashed var(--ink);padding-top:12px;
  font-size:11.5px;color:var(--muted)}}
footer .mono{{word-break:break-all;color:var(--ink)}}
@media print{{body{{padding:0}}.stamp{{transform:rotate(-4deg)}}}}
</style></head><body>
<header>
  <div>
    <div class="eyebrow">{_esc(report.get('bank_name') or 'PrefectOS')} · Governed loan document processing</div>
    <h1>{_esc(doc_title)} — {_esc(job.loan_label)} · <span class="mono">{_esc(applicant)}</span></h1>
    <div class="small muted">Generated {_fmt_dt(report.get('generated_at', ''))}
      · Job <span class="mono">{_esc(job.job_id)}</span>
      · Processed in {_esc(fmt_duration(report.get('elapsed_s')))}
      · {_esc(fmt_cost(report.get('cost_usd')))}</div>
  </div>
  <div class="stamp">{_esc(stamp_text)}<small>{_esc(stamp_sub)}</small></div>
</header>

<h2>Applicant file</h2>
<div class="fields">
  <div class="f"><b>Loan product</b>{_esc(job.loan_label)}</div>
  <div class="f"><b>Applicant</b>{_esc(applicant)}</div>
  <div class="f"><b>Confidence</b>{_esc(a["confidence"] or "—")}</div>
  <div class="f"><b>Documents</b><span class="mono">{report['documents_processed']}</span> processed
    · <span class="mono">{report['documents_failed']}</span> unreadable</div>
  <div class="f"><b>Agents</b><span class="mono">{len(report.get('agents') or [])}</span> ·
    {_esc(', '.join(ag['label'] for ag in (report.get('agents') or [])))}</div>
  <div class="f"><b>Tokens</b><span class="mono">{job.tokens_in:,}</span> in ·
    <span class="mono">{job.tokens_out:,}</span> out</div>
  <div class="f"><b>Cost (list price)</b><span class="mono">{_esc(fmt_cost(report.get('cost_usd')))}</span>
    · {_esc(report.get('model', ''))}</div>
  <div class="f"><b>Processing time</b><span class="mono">{_esc(fmt_duration(report.get('elapsed_s')))}</span></div>
  <div class="f"><b>AI share</b><span class="mono">{report.get('ai_share', 1.0) * 100:.0f}%</span>
    · <span class="mono">{report.get('documents_reconciled_in_code', 0)}</span> reconciled in code
    · <span class="mono">{report.get('documents_escalated_to_ai', 0)}</span> escalated</div>
  <div class="f"><b>Per document</b><span class="mono">{_esc(fmt_duration(
      (report.get('elapsed_s') or 0) / max(len(job.docs), 1)))}</span> average</div>
  <div class="f" style="grid-column:1/-1;border-bottom:none"><b>Input path</b>
    <span class="mono small">{_esc(job.input_path)}</span></div>
</div>

<h2>{_esc(verdict_word)}</h2>
<div class="figs">
  <div><b>{_esc(verdict_word)}</b><span class="mono">{_esc(decision)}</span></div>
  <div><b>Checks met</b><span class="mono">{sum(
      1 for c in (a["checks"]) if isinstance(c, dict)
      and c.get('status') == 'met')}/{len(a["checks"])}</span></div>
  <div><b>Missing documents</b><span class="mono">{len(a["missing_documents"])}</span></div>
  <div><b>Risk flags</b><span class="mono">{len(a["risk_flags"])}</span></div>
</div>
<div class="rationale">{_esc(a["rationale"]
    or report.get("assessment_text") or 'No rationale was returned.')}</div>

{findings_block}
{crit_block}
{list_block("Missing documents", a["missing_documents"])}
{list_block("Risk flags", a["risk_flags"])}
{list_block("Next steps", a["next_steps"])}

<h2>Documents processed ({len(job.docs)})</h2>
<table class="grid"><thead><tr><th>Document</th><th>Type</th><th>Relevance</th>
<th>Key facts</th><th style="text-align:right">Tokens</th></tr></thead>
<tbody>{doc_rows}</tbody></table>

{plan_block}

{policy_block}

<h2>Agents ({len(report.get('agents') or [])} of {MAX_AGENTS_PER_JOB})</h2>
<table class="grid"><thead><tr><th>Agent</th><th>Role</th><th>Role card</th><th>Model</th>
<th style="text-align:right">Calls</th><th style="text-align:right">Tokens in</th>
<th style="text-align:right">Tokens out</th><th style="text-align:right">Cost</th></tr></thead>
<tbody>{agent_rows}</tbody></table>

<footer>
  <b>{_esc(report.get('bank_name') or 'PrefectOS')}</b> — {_esc(job.loan_label)} {_esc(doc_title.lower())}, generated {_fmt_dt(report.get('generated_at', ''))}. For internal credit
  review only.<br>
  Report SHA-256: <span class="mono">{payload_sha}</span><br>
  Written to <span class="mono">{_esc(job.run_dir or job.output_path)}</span> alongside
  <span class="mono">eligibility_report.json</span>,
  <span class="mono">processing_plan.json</span> and the per-document findings under
  <span class="mono">documents/</span>. This report is generated deterministically from
  the sealed JSON — regenerate it at any time from that file.
</footer>
</body></html>"""
