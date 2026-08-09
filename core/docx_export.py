# Auto-split from Orchestrator.py — part of the PrefectOS core package.
"""Markdown → Word export for approved plan/spec documents."""
from __future__ import annotations

import re
from pathlib import Path

from .config import log

# ─────────────────────────────────────────────────────────────────────────────

def _docx_add_runs(paragraph, text: str) -> None:
    """Add text to a paragraph, honouring **bold** and `inline code`."""
    for token in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(token)


def export_docx(title: str, md_text: str, dest: Path) -> Path | None:
    """Convert markdown to a Word document (headings, bullet/numbered lists,
    fenced code blocks, bold/inline code). Best-effort: returns None and logs
    a warning if python-docx is not installed."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        log.warning("python-docx not installed — skipping %s (pip install python-docx)", dest.name)
        return None

    doc = Document()
    doc.add_heading(title, level=0)

    in_code = False
    code_lines: list[str] = []
    for line in md_text.splitlines():
        if line.strip().startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run  = para.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_lines = []
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)", line)
        bullet  = re.match(r"^\s*[-*]\s+(.*)", line)
        number  = re.match(r"^\s*\d+[.)]\s+(.*)", line)
        if heading:
            doc.add_heading(heading.group(2).strip(), level=min(len(heading.group(1)), 4))
        elif bullet:
            _docx_add_runs(doc.add_paragraph(style="List Bullet"), bullet.group(1))
        elif number:
            _docx_add_runs(doc.add_paragraph(style="List Number"), number.group(1))
        elif line.strip():
            _docx_add_runs(doc.add_paragraph(), line.strip())

    doc.save(str(dest))
    log.info("Word export: %s", dest)
    return dest


